import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

doc = Document()

# Configure Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Typography Configuration
styles = doc.styles
normal_style = styles['Normal']
normal_style.font.name = 'Arial'
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(30, 41, 59)

# ------------------------------------------------------------------------------
# DOCUMENT HEADER / TITLE
# ------------------------------------------------------------------------------
title = doc.add_paragraph()
title_run = title.add_run("Project 5 Oral Defense Dossier\n")
title_run.font.name = 'Georgia'
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(15, 23, 42)

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run("Production Blueprint: Explainable & Adversarially Robust IDS\n")
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(14)
sub_run.font.color.rgb = RGBColor(245, 158, 11)
sub_run.font.bold = True

# Separator Line
p_sep = doc.add_paragraph()
p_sep_run = p_sep.add_run("_________________________________________________________________________________")
p_sep_run.font.color.rgb = RGBColor(203, 213, 225)

# ==============================================================================
# SECTION 1: SLIDE-BY-SLIDE SCRIPT
# ==============================================================================
h1 = doc.add_paragraph()
h1_run = h1.add_run("Part 1: Comprehensive Oral Defense Script")
h1_run.font.name = 'Georgia'
h1_run.font.size = Pt(18)
h1_run.font.bold = True
h1_run.font.color.rgb = RGBColor(15, 23, 42)
h1.space_before = Pt(18)
h1.space_after = Pt(12)

slides_data = [
    ("Slide 1: Title & Core Hypothesis",
     "Good morning, distinguished members of the jury. Today I will present our work on Project 5: Building an Explainable and Adversarially Robust Intrusion Detection System. In traditional machine learning security architectures, explainability and adversarial robustness are frequently treated as distinct, secondary engineering objectives. Our core hypothesis directly challenges this siloed approach: we argue that in adversarial environments, model explanations are inherently dual-use. While they empower network security analysts to validate anomalous alerts and accelerate incident response, they simultaneously provide an adaptive adversary with a deterministic map of feature attributions that can be leveraged to guide highly effective evasion attacks. This presentation details the design, implementation, and rigorous validation of an Adv+ExtraTrees Ensemble IDS that optimizes clean performance while systematically minimizing adversarial transferability.\n\nTRANSITION: With this fundamental structural tension between transparency and vulnerability established, let us examine the specific engineering objectives we set out to meet."),
    
    ("Slide 2: Project Objective and Requirements",
     "Our project was structured around five core engineering requirements: establishing high-fidelity baseline classifiers, applying mathematically diverse explainability frameworks, quantifying explanation stability under resampling, evaluating explanation-guided adversarial evasion vectors, and finally, delivering a fully reproducible, hardened pipeline. Crucially, we do not treat explainability as a post-hoc diagnostic luxury. Instead, we treat explanation stability and adversarial evasion rates as core security metrics. By evaluating both trust and threat simultaneously, we ensure that our final system remains resilient even when its internal classification logic is partially or fully exposed to a threat actor.\n\nTRANSITION: To ground these high-level objectives in a realistic security context, we utilized a standard benchmark dataset that highlights specific engineering challenges regarding class imbalance."),
    
    ("Slide 3: Dataset and Preprocessing Pipeline",
     "We executed our evaluation on the NSL-KDD dataset, partitioning our pipeline across the standard 125,973 training records in KDDTrain+ and 22,544 records in KDDTest+. Our preprocessing workflow was strictly engineered to completely eliminate data leakage: categorical features—specifically protocol_type, service, and flag—were one-hot encoded, and MinMax scaling parameters were fitted exclusively on the training distribution before being applied to the test split. Incorporating domain-specific feature engineering—including byte logs, traffic totals, byte ratios, service ratios, error-rate gaps, and login anomaly scores—expanded our final representation space to 478 dimensions. Operationally, we isolated Remote-to-Local (R2L) and User-to-Root (U2R) attack families as our primary pain points. These rare attack vectors are heavily masked by standard global accuracy metrics, demanding an explicit focus on per-family recall.\n\nTRANSITION: To bridge the gap between these heavily skewed feature spaces and a highly resilient deployment posture, we constructed a multi-stage model architecture."),
    
    ("Slide 4: Model Pipeline Architecture",
     "Our architectural pipeline scales systematically from linear baselines to a robust heterogeneous ensemble. We began with a Logistic Regression baseline, progressing to a Random Forest paired with TreeSHAP, and a PyTorch Multi-Layer Perceptron utilizing Integrated Gradients and SmoothIG. To actively defend this infrastructure, the neural architecture was subjected to multi-epsilon Projected Gradient Descent (PGD) adversarial fine-tuning, while the tree infrastructure was hardened via SHAP-guided robust training. Our final production system is a soft-voting ensemble combining the adversarially trained Torch Binary MLP with the SHAP-Robust ExtraTrees classifier. This specific design pairs the non-linear gradient diversity of the deep neural network with the rigid, axis-aligned decision boundaries of the tree ensemble, intentionally disrupting uniform optimization landscapes for attackers.\n\nTRANSITION: Let us now evaluate how this structural synergy manifests in clean performance metrics under rigorous operational conditions."),
    
    ("Slide 5: Main Clean Results & Threshold Optimization",
     "When analyzing clean network traffic, our Adv+ExtraTrees Ensemble outpaced all standalone baselines, yielding a binary precision of 0.9062, a binary recall of 0.9064, an F1 score of 0.9063, and a PR-AUC of 0.9626. A critical engineering decision was adjusting our classification threshold downward to 0.15. In actual network deployment settings, maintaining a standard 0.5 decision threshold introduces unacceptable operational risk by missing low-volume, high-severity compromises. Furthermore, while our optional 'Security-OR' configuration demonstrates maximum adversarial resistance, it suppresses our clean F1 score to 0.8796 due to an elevated false positive rate, justifying our selection of the Ensemble as the primary model.\n\nTRANSITION: The true validation of this threshold tuning, however, becomes apparent when we break down our detection capabilities across individual attack families."),
    
    ("Slide 6: Rare-Family Recall Verification",
     "In real-world network security, a high global F1 score can easily mask a catastrophic vulnerability if the underlying model fails to detect low-frequency intrusion attempts. As displayed in this data breakdown, our model maintains exceptional detection on high-volume vectors, yielding a DoS recall of 0.9859 and a Probe recall of 1.0000. Crucially, for the highly covert and dangerous R2L and U2R vectors, our model reaches a recall of 0.6786 and 0.8060, respectively. While these minority classes remain inherently challenging due to extreme data scarcity in training and severe test-set distribution shifts, our ensemble significantly outperforms traditional standalone deep baselines.\n\nTRANSITION: Having established the system's clean data efficacy, we now pivot to our second core requirement: dissecting the explainability frameworks that audit these decisions."),
    
    ("Slide 7: Complementary Explainability Methods",
     "To ensure comprehensive auditing capability, we deployed three complementary explainability methods targeting distinct algorithmic paradigms. We utilized TreeSHAP to compute exact feature attributions for our tree structures, Integrated Gradients to satisfy the axiom of completeness across our neural pathways, and SmoothIG to eliminate local gradient noise by averaging attributions over a Gaussian input neighborhood. Our final ensemble explanation normalizes and aggregates these outputs. The dominant signals extracted are highly interpretable from a network security perspective: connection flag states, same-service consistency rates, host-service concentration levels, and error-rate anomalies emerge as the primary pillars of our detection logic.\n\nTRANSITION: However, the core security dilemma of our project arises when we evaluate the stability of these identical explanations."),
    
    ("Slide 8: Quantifying the Stability-Risk Trade-off",
     "We evaluated explanation stability using both Local and Bootstrap Jaccard metrics to track top-feature consistency under data resampling. While the Random Forest SHAP baseline drops to a Bootstrap Jaccard of 0.5744 under resampling, our neural SmoothIG framework achieves a highly stable 0.9564, and our final ensemble maintains a robust global Jaccard of 0.8873. This high consistency is incredibly beneficial for establishing security analyst trust. Yet, from an offensive standpoint, this stability introduces a severe dual-use risk. If an explanation is highly deterministic across diverse data distributions, it signals to an attacker that the model's feature reliance is static and predictable, making it a reliable target for crafting targeted evasion vectors.\n\nTRANSITION: To verify this threat model empirically, we executed explanation-guided attacks directly against our infrastructure."),
    
    ("Slide 9: Adversarial Attack Vulnerability Analysis",
     "Our vulnerability testing confirms that interpretability can be actively weaponized. When an adversary exploits standalone TreeSHAP or Integrated Gradients attributions, they successfully achieve evasion rates of 18.22% against the Random Forest and 15.97% against the standalone Torch MLP using minimal feature modifications. Furthermore, unconstrained multi-step white-box PGD completely compromises the unhardened architectures, achieving 100.00% evasion across all testing scales. Even after adversarial fine-tuning, direct white-box deep PGD remains a formidable threat at maximum epsilon values, which underscores why standalone deep models are structurally unsuited for high-threat deployments.\n\nTRANSITION: This limitation highlights the core contribution of our defense results: leveraging cross-model architectural asymmetry to break adversarial transferability."),
    
    ("Slide 10: Main Defense Results: Eliminating Transfer-PGD",
     "Our most compelling security finding lies in the systematic elimination of adversarial transferability. In this threat matrix, adversarial examples are optimally generated against a fully exposed neural surrogate model and subsequently transferred to our target ensemble. As the data demonstrates, while the surrogate model is heavily compromised—scaling from a 37.29% evasion rate at epsilon 0.03 to an absolute 100.00% at epsilon 0.15—our Adv+ExtraTrees Ensemble systematically neutralizes the attack. It suppresses transfer evasion to a mere 0.50% at epsilon 0.03, and achieves an absolute 0.00% evasion rate across all higher budgets. We explicitly characterize this as empirical robustness bounded by our tested transfer threat model, rather than an absolute mathematical certification.\n\nTRANSITION: To maintain academic rigor, we must look closely at the engineering limitations that bound these claims before concluding."),
    
    ("Slide 11: Engineering and Deployment Limitations",
     "We acknowledge four core limitations in this implementation. First, the NSL-KDD benchmark contains historical traffic patterns that do not capture the semantic complexity of modern, zero-day cloud exploits. Second, our adversarial evaluation occurs entirely in feature-space, meaning we manipulate tabular dimensions directly. While mutable-feature PGD introduces realistic bounds, feature-space optimization does not guarantee packet-space validity, as it ignores protocol constraints like TCP handshakes and checksum calculations. Finally, the computational overhead of calculating real-time SHAP values represents a clear bottleneck for multi-gigabit line-rate network deployments, requiring asynchronous or out-of-band architectural positioning.\n\nTRANSITION: In conclusion, let us summarize the core architectural contributions of this framework."),
    
    ("Slide 12: Conclusion and Strategic Takeaways",
     "In conclusion, Project 5 successfully validates an interpretable, hardened intrusion detection framework that safely navigates the trade-off between transparency and security. By combining an adversarially fine-tuned network core with a SHAP-Robust ExtraTrees model, we achieve an optimal balance: high clean data fidelity with an F1 score of 0.9063, strong rare-attack visibility, and the near-total suppression of transfer-based adversarial evasion. This architecture establishes that deep learning systems can be deployed transparently in adversarial settings, provided that cross-model architectural asymmetries are deliberately engineered into the defense pipeline. Thank you, and I am now open to your questions.")
]

for title_text, script_text in slides_data:
    p_t = doc.add_paragraph()
    r_t = p_t.add_run(title_text)
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(13)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(245, 158, 11)
    p_t.space_before = Pt(12)
    p_t.space_after = Pt(4)
    
    p_s = doc.add_paragraph()
    r_s = p_s.add_run(script_text)
    p_s.paragraph_format.line_spacing = 1.15
    p_s.space_after = Pt(12)

# ==============================================================================
# SECTION 2: HIGH PERFORMANCE Q&A
# ==============================================================================
doc.add_page_break()
h2 = doc.add_paragraph()
h2_run = h2.add_run("Part 2: High-Performance Q&A Defense Matrix")
h2_run.font.name = 'Georgia'
h2_run.font.size = Pt(18)
h2_run.font.bold = True
h2_run.font.color.rgb = RGBColor(15, 23, 42)
h2.space_before = Pt(18)
h2.space_after = Pt(12)

qa_data = [
    ("Q1: Your final model chooses an operational threshold of 0.15. Doesn't lowering the threshold from 0.5 drastically increase your False Positive Rate on clean traffic? How do you justify this operationally?",
     "Yes, lowering the classification threshold to 0.15 does increase the clean False Positive Rate, which is why our global clean F1 score sits at 0.9063 instead of a higher value. However, from a practical security engineering perspective, this trade-off is entirely justified. In network intrusion detection, the cost functions of operational errors are highly asymmetric. A False Positive results in an analyst auditing a benign alert, which can be handled by automated downstream filtering or tier-1 triage workflows. A False Negative, particularly for highly destructive, low-frequency attack families like U2R or R2L, can mean a silent, full system compromise. By minimizing the threshold to 0.15, we optimize our balanced accuracy to 0.9064 and specifically pull our U2R recall up to 0.8060, creating a defensive posture appropriate for a real-world Security Operations Center."),
    
    ("Q2: You emphasize that explanation stability poses a 'dual-use risk.' Mathematically or algorithmically, why does a stable explanation make a model more vulnerable to an attacker?",
     "Explanation stability means that across bootstrap resamples or local input perturbations, the top feature attributions remain consistent—as seen in our SmoothIG bootstrap stability score of 0.9564. Algorithmic vulnerability occurs because a stable explanation proves that the model's local decision boundary is highly deterministic, continuous, and dependent on a static subset of features. If an attacker queries a sample and receives an explanation that is highly stable, they know that the gradient direction or tree path attribution they have discovered is not an artifact of local data noise. It gives them a high-fidelity map indicating exactly which features to modify (e.g., altering same_srv_rate or logged_in markers) to cross the decision boundary with minimal perturbation, maximizing their attack success while remaining covert."),
    
    ("Q3: What is the fundamental mathematical distinction between TreeSHAP and Integrated Gradients? Why use both in an ensemble?",
     "The mathematical foundations of the two methods are distinct, allowing them to capture entirely different structural dependencies. TreeSHAP is designed for tree structures; it calculates exact Shapley values by leveraging the tree topology to compute conditional expectations across all feature subsets in polynomial time, ignoring features that are not on the active splitting paths. Integrated Gradients is a gradient-based method designed for continuous neural networks. It computes the path integral of the gradients along a straight line from a defined baseline input to the target instance, satisfying the Axiom of Completeness, meaning the attributions sum precisely to the difference between the target output and baseline output. By combining them in our ensemble explanation, we capture both the sharp, axis-aligned step attributions of our ExtraTrees component and the smooth, continuous gradient trajectories of our PyTorch MLP, preventing an adversary from optimizing an attack against a single mathematical formulation."),
    
    ("Q4: Your results show that the Adv+ExtraTrees Ensemble reduced transfer-PGD evasion to 0.00% at epsilon 0.06 and above. Does this mean your system is completely robust against adversarial attacks?",
     "Absolutely not, and we explicitly avoid making that claim to maintain academic rigor. Our model demonstrates high empirical robustness under the specific transfer-PGD threat model tested, where the adversary generates white-box attacks against a neural surrogate and transfers them to our ensemble. The reason evasion drops to 0.00% is that the tree component—the SHAP-Robust ExtraTrees—does not rely on continuous gradients. Therefore, adversarial examples optimized to exploit the continuous gradient landscape of the PyTorch MLP fail to transfer effectively across the non-continuous, axis-aligned decision boundaries of the tree model. However, an adaptive adversary executing a direct white-box attack against the ensemble itself—using zeroth-order optimization or black-box genetic algorithms—could potentially identify evasion vectors. Our robustness is empirical and bounded, not mathematically certified."),
    
    ("Q5: You noted that your adversarial optimization occurs in 'feature-space' rather than 'packet-space.' If you were to transition this project from a laboratory setting to a live enterprise network interface card (NIC), what engineering breakdowns would occur?",
     "In a feature-space evaluation, the optimization algorithm treats the tabular inputs as an arbitrary vector of independent floats, allowing it to modify fields at will. If we deploy this model directly onto a live network interface card, two major engineering breakdowns would occur:\n1. Protocol and Semantic Invalidation: The feature-space attack might decrease a feature like src_bytes while simultaneously increasing a feature like duration in a way that violates the physics of network protocols. In packet-space, you cannot arbitrarily change byte counts without modifying actual payload lengths, recomputing TCP checksums, adjusting sequence numbers, and maintaining stateful handshakes. Most of our feature-space adversarial samples would be dropped by the OS network stack for being malformed before they even reached the application layer.\n2. Computational Infeasibility: Extracting our 478 features—such as calculating historical same-service rates over moving time windows—requires a stateful parsing engine. Doing this at a line-rate of 10 or 40 Gigabits per second, while simultaneously computing real-time TreeSHAP expectations, would introduce severe packet processing latency and memory exhaustion, resulting in dropped packets and unmonitored traffic.")
]

for q, a in qa_data:
    p_q = doc.add_paragraph()
    r_q = p_q.add_run(q)
    r_q.font.name = 'Arial'
    r_q.font.size = Pt(12)
    r_q.font.bold = True
    r_q.font.color.rgb = RGBColor(15, 23, 42)
    p_q.space_before = Pt(10)
    p_q.space_after = Pt(4)
    
    p_a = doc.add_paragraph()
    r_a = p_a.add_run(a)
    p_a.paragraph_format.line_spacing = 1.15
    p_a.space_after = Pt(10)

# ==============================================================================
# SECTION 3: DEFINITIVE BACKUP METRICS
# ==============================================================================
doc.add_page_break()
h3 = doc.add_paragraph()
h3_run = h3.add_run("Part 3: Definitive Performance Metrics Matrix")
h3_run.font.name = 'Georgia'
h3_run.font.size = Pt(18)
h3_run.font.bold = True
h3_run.font.color.rgb = RGBColor(15, 23, 42)
h3.space_before = Pt(18)
h3.space_after = Pt(12)

metrics_table_data = [
    ("Primary Selected Production Architecture", "Adv+ExtraTrees Ensemble IDS"),
    ("Operational Decision Threshold", "0.15"),
    ("Global Binary Precision", "0.9062"),
    ("Global Binary Recall", "0.9064"),
    ("Global Binary F1-Score", "0.9063"),
    ("Precision-Recall Area Under Curve (PR-AUC)", "0.9626"),
    ("Balanced Accuracy Score", "0.9064"),
    ("Denial of Service (DoS) Family Recall", "0.9859"),
    ("Probe / Scanning Family Recall", "1.0000"),
    ("Remote-to-Local (R2L) Family Recall", "0.6786"),
    ("User-to-Root (U2R) Family Recall", "0.8060"),
    ("Neural SmoothIG Bootstrap Jaccard Stability", "0.9564"),
    ("Ensemble Aggregated Explanation Stability", "0.8873"),
    ("Transfer-PGD Evasion Rate (epsilon = 0.03)", "0.50%"),
    ("Transfer-PGD Evasion Rate (epsilon >= 0.06)", "0.00% (Complete Suppression)"),
    ("Optional Strict Mode (Security-OR) Clean F1", "0.8796")
]

t = doc.add_table(rows=1, cols=2)
t.autofit = False
t.columns[0].width = Inches(4.5)
t.columns[1].width = Inches(2.5)

hdr_cells = t.rows[0].cells
hdr_cells[0].text = 'Metric Dimension / Operational Item'
hdr_cells[1].text = 'Value / Configuration'
for cell in hdr_cells:
    cell.fill.solid()
    cell.fill.fore_color.rgb = RGBColor(30, 41, 59)
    for p in cell.text_frame.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)

for metric, val in metrics_table_data:
    row_cells = t.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = val
    
    # Shade specific rows or add formatting
    if "Ensemble IDS" in val or "0.00%" in val:
        for cell in row_cells:
            cell.fill.solid()
            cell.fill.fore_color.rgb = RGBColor(241, 245, 249)

for row in t.rows:
    for cell in row.cells:
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        for p in cell.text_frame.paragraphs:
            p.space_after = Pt(2)
            p.space_before = Pt(2)
            for r in p.runs:
                r.font.size = Pt(10)

doc.save("Project5_Oral_Script_and_QA.docx")
print("Document successfully compiled and saved as 'Project5_Oral_Script_and_QA.docx'.")